import os
import json
import threading
import time
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
from docx import Document
import docx
import pysrt
import google.generativeai as genai
from datetime import datetime
import re

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.secret_key = 'supersecretkey'

# ایجاد دایرکتوری‌ها
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# تنظیمات ترجمه
RTL_LANGUAGES = ['fa', 'ar', 'he']  # زبان‌های راست به چپ
TRANSLATION_STYLES = {
    'standard': 'ترجمه استاندارد و رسمی',
    'colloquial': 'ترجمه محاوره‌ای و غیررسمی',
    'literary': 'ترجمه ادبی و رمان',
    'technical': 'ترجمه فنی و تخصصی'
}

# وضعیت پیشرفت ترجمه
translation_progress = {}
translation_lock = threading.Lock()

def process_document(file_path, api_key, seed, style, source_lang, target_lang, job_id):
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            process_docx(file_path, api_key, seed, style, source_lang, target_lang, job_id)
        elif file_ext == '.srt':
            process_srt(file_path, api_key, seed, style, source_lang, target_lang, job_id)
        else:
            with translation_lock:
                translation_progress[job_id] = {
                    'status': 'error',
                    'message': 'فرمت فایل پشتیبانی نمی‌شود'
                }
    except Exception as e:
        with translation_lock:
            translation_progress[job_id] = {
                'status': 'error',
                'message': str(e)
            }

def process_docx(file_path, api_key, seed, style, source_lang, target_lang, job_id):
    doc = Document(file_path)
    total_elements = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
    processed_elements = 0
    
    # پیکربندی Gemini
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    # ترجمه پاراگراف‌ها
    for para in doc.paragraphs:
        if para.text.strip():
            translated = translate_text(
                para.text, 
                model, 
                style, 
                source_lang, 
                target_lang,
                seed
            )
            para.text = translated
            
            # تنظیم تراز راست برای زبان‌های RTL
            if target_lang in RTL_LANGUAGES:
                para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        
        processed_elements += 1
        update_progress(job_id, processed_elements, total_elements)
    
    # ترجمه جداول
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translated = translate_text(
                            para.text, 
                            model, 
                            style, 
                            source_lang, 
                            target_lang,
                            seed
                        )
                        para.text = translated
                        
                        # تنظیم تراز راست برای زبان‌های RTL
                        if target_lang in RTL_LANGUAGES:
                            para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                
                processed_elements += 1
                update_progress(job_id, processed_elements, total_elements)
    
    # ذخیره سند ترجمه شده
    output_path = os.path.join(
        app.config['OUTPUT_FOLDER'],
        f"translated_{secure_filename(os.path.basename(file_path))}"
    )
    doc.save(output_path)
    
    with translation_lock:
        translation_progress[job_id] = {
            'status': 'completed',
            'output': output_path,
            'progress': 100
        }

def process_srt(file_path, api_key, seed, style, source_lang, target_lang, job_id):
    subs = pysrt.open(file_path)
    total_subs = len(subs)
    
    # پیکربندی Gemini
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    # ترجمه زیرنویس‌ها
    for i, sub in enumerate(subs):
        if sub.text.strip():
            translated = translate_text(
                sub.text, 
                model, 
                style, 
                source_lang, 
                target_lang,
                seed
            )
            sub.text = translated
        
        update_progress(job_id, i + 1, total_subs)
    
    # ذخیره فایل ترجمه شده
    output_path = os.path.join(
        app.config['OUTPUT_FOLDER'],
        f"translated_{secure_filename(os.path.basename(file_path))}"
    )
    subs.save(output_path, encoding='utf-8')
    
    with translation_lock:
        translation_progress[job_id] = {
            'status': 'completed',
            'output': output_path,
            'progress': 100
        }

def translate_text(text, model, style, source_lang, target_lang, seed=None):
    try:
        # ایجاد دستورالعمل ترجمه با توجه به سبک انتخابی
        style_instruction = TRANSLATION_STYLES.get(style, 'ترجمه استاندارد')
        
        prompt = f"""
        شما یک مترجم حرفه‌ای هستید. متن زیر را از زبان {source_lang} به {target_lang} با سبک {style_instruction} ترجمه کنید.
        قوانین:
        1. فقط ترجمه نهایی را برگردانید (بدون توضیح اضافه)
        2. ساختار و معنی اصلی را کاملاً حفظ کنید
        3. برای اصطلاحات تخصصی دقت کنید
        4. متن ترجمه شده باید روان و طبیعی باشد
        
        متن برای ترجمه:
        "{text}"
        """
        
        # تنظیمات تولید
        generation_config = {
            "temperature": 0.3,
            "top_p": 1,
            "top_k": 32,
            "max_output_tokens": 2000,
        }
        
        # استفاده از seed اگر وجود دارد
        if seed:
            generation_config["seed"] = int(seed)
        
        response = model.generate_content(
            prompt,
            generation_config=generation_config
        )
        
        return response.text.strip('"')
    
    except Exception as e:
        print(f"Error in translation: {str(e)}")
        return text  # در صورت خطا متن اصلی را برگردان

def update_progress(job_id, current, total):
    progress = int((current / total) * 100) if total > 0 else 0
    with translation_lock:
        if job_id in translation_progress:
            translation_progress[job_id]['progress'] = progress
        else:
            translation_progress[job_id] = {
                'status': 'processing',
                'progress': progress
            }

@app.route('/')
def index():
    return render_template(
        'index.html',
        styles=TRANSLATION_STYLES,
        rtl_langs=RTL_LANGUAGES
    )

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    # پارامترهای فرم
    api_key = request.form.get('api_key', '')
    seed = request.form.get('seed', '')
    style = request.form.get('style', 'standard')
    source_lang = request.form.get('source_lang', 'en')
    target_lang = request.form.get('target_lang', 'fa')
    
    if not api_key:
        return jsonify({'error': 'API Key is required'}), 400
    
    # ذخیره فایل آپلود شده
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    # ایجاد ID کار
    job_id = f"job_{int(time.time())}_{os.urandom(4).hex()}"
    
    # شروع ترجمه در نخ جداگانه
    thread = threading.Thread(
        target=process_document,
        args=(file_path, api_key, seed, style, source_lang, target_lang, job_id)
    )
    thread.start()
    
    # وضعیت اولیه
    with translation_lock:
        translation_progress[job_id] = {
            'status': 'started',
            'progress': 0
        }
    
    return jsonify({'job_id': job_id})

@app.route('/progress/<job_id>')
def get_progress(job_id):
    with translation_lock:
        progress_data = translation_progress.get(job_id, {'status': 'not_found'})
    return jsonify(progress_data)

@app.route('/download/<job_id>')
def download_file(job_id):
    with translation_lock:
        job_data = translation_progress.get(job_id, {})
    
    if job_data.get('status') == 'completed' and 'output' in job_data:
        return send_file(
            job_data['output'],
            as_attachment=True,
            download_name=os.path.basename(job_data['output'])
        )
    
    return jsonify({'error': 'File not ready or not found'}), 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
